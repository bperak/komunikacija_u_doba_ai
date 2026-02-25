"""
Izvlačenje Mermaid blokova iz DOCX rukopisa.
Blokovi počinju s ```mermaid ili ```, završavaju s ``` ili end```.
Izlaz: lista Mermaid nizova + opcionalno manuscript/mermaid_diagrami.md
"""

from __future__ import annotations

import json
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

# Word OOXML namespaces
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _text_of_elt(elt: ET.Element, ns: str) -> str:
    out = [elt.text or ""]
    for c in elt:
        out.append(_text_of_elt(c, ns))
        if c.tail:
            out.append(c.tail)
    return "".join(out)


def get_paragraphs_from_docx_via_zip(docx_path: Path) -> list[str]:
    """Izvadi paragrafe iz docx kao listu nizova."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    parts = []
    with zipfile.ZipFile(docx_path, "r") as z:
        with z.open("word/document.xml") as f:
            tree = ET.parse(f)
            root = tree.getroot()
            for p in root.iter("{%s}p" % ns):
                parts.append(_text_of_elt(p, ns))
    return parts


def get_full_text_from_docx_via_zip(docx_path: Path) -> str:
    """Izvadi tekst iz docx kao zip (word/document.xml), paragrafi odvojeni \\n."""
    return "\n".join(get_paragraphs_from_docx_via_zip(docx_path))


def get_paragraphs_from_docx(docx_path: str | Path) -> list[str]:
    """Vraća listu paragrafa (prvo zip, pa python-docx)."""
    path = Path(docx_path) if isinstance(docx_path, str) else docx_path
    try:
        return get_paragraphs_from_docx_via_zip(path)
    except (zipfile.BadZipFile, KeyError, ET.ParseError):
        pass
    from docx import Document
    with open(path, "rb") as f:
        doc = Document(f)
    return [p.text or "" for p in doc.paragraphs]


def get_full_text_from_docx(docx_path: str | Path) -> str:
    """Složi cijeli tekst dokumenta (prvo pokušaj zip, pa python-docx)."""
    path = Path(docx_path) if isinstance(docx_path, str) else docx_path
    try:
        return get_full_text_from_docx_via_zip(path)
    except (zipfile.BadZipFile, KeyError, ET.ParseError):
        pass
    from docx import Document
    with open(path, "rb") as f:
        doc = Document(f)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_mermaid_blocks(full_text: str) -> list[str]:
    """
    Pronađi sve blokove koji počinju s ```mermaid ili ``` i završavaju s ```.
    Ako blok završava s 'end' pa ```, 'end' se ne uključuje u kod.
    Vraća listu nizova (samo Mermaid izvorni kod, bez ogradnih ```).
    """
    blocks = []
    # Tri backticka kao literal
    fence = "```"
    # Mogući početak: ```mermaid ili ```
    start_pattern = re.compile(r"```(?:mermaid)?\s*\n", re.IGNORECASE)
    text = full_text

    while True:
        match = start_pattern.search(text)
        if not match:
            break
        start_end = match.end()
        # Odmah nakon početka tražimo kraj: ``` ili end pa ```
        rest = text[start_end:]
        # Kraj 1: samo ```
        close_plain = rest.find("\n" + fence)
        # Kraj 2: end pa newline pa ```
        close_with_end = re.search(r"\nend\s*\n\s*" + re.escape(fence), rest, re.IGNORECASE)
        if close_with_end:
            end_pos = close_with_end.start()
            code = rest[:end_pos].strip()
            # Ukloni eventualni trailing "end" na kraju zadnjeg reda
            if code.endswith("end"):
                code = code[:-3].strip()
        elif close_plain != -1:
            code = rest[:close_plain].strip()
        else:
            # Nema zatvaranja - uzmi do kraja teksta (fallback)
            code = rest.strip()
        code = code.strip()
        if not code:
            pass
        elif code.startswith("graph ") or (code.split("\n")[0].strip().startswith("%%") and "graph " in code):
            blocks.append(code)
        # Sljedeći blok: pomakni se na prvi zatvarajući ```
        next_close = rest.find(fence)
        if next_close == -1:
            break
        text = rest[next_close + len(fence) :]
    return blocks


def clean_caption(s: str) -> str:
    """
    Čisti caption iz DOCX-a: uklanja broj poglavlja na početku, broj stranice na kraju,
    asteriske (kurziv), normalizira razmake. Za bolje mapiranje na MD.
    """
    if not s:
        return ""
    s = (s or "").strip()
    s = re.sub(r"^#+\s*", "", s)
    s = re.sub(r"^\s*\d+(?:\.\d+)*\.?\s*", "", s)  # 2.2.2 ili 2.7.2 na početku
    s = re.sub(r"\s*\d{1,4}\s*$", "", s)  # broj stranice na kraju (npr. 13, 26)
    s = s.strip("*").strip()
    s = re.sub(r"\s+", " ", s).strip()
    return s


def count_all_fenced_blocks(full_text: str) -> tuple[int, int]:
    """
    Broji sve ``` blokove u tekstu i koliko njih sadrži 'graph ' (mermaid).
    Vraća (ukupno_blokova, blokova_s_graph).
    """
    fence = "```"
    start_pattern = re.compile(r"```(?:mermaid)?\s*\n", re.IGNORECASE)
    text = full_text
    total = 0
    with_graph = 0
    while True:
        match = start_pattern.search(text)
        if not match:
            break
        start_end = match.end()
        rest = text[start_end:]
        close_plain = rest.find("\n" + fence)
        close_with_end = re.search(r"\nend\s*\n\s*" + re.escape(fence), rest, re.IGNORECASE)
        if close_with_end:
            code = rest[: close_with_end.start()].strip()
            if code.endswith("end"):
                code = code[:-3].strip()
        elif close_plain != -1:
            code = rest[:close_plain].strip()
        else:
            code = rest.strip()
        code = code.strip()
        total += 1
        if code and (code.startswith("graph ") or (code.split("\n")[0].strip().startswith("%%") and "graph " in code)):
            with_graph += 1
        next_close = rest.find(fence)
        if next_close == -1:
            break
        text = rest[next_close + len(fence) :]
    return total, with_graph


def extract_blocks_with_captions(full_text: str, paragraphs: list[str], only_mermaid: bool = True) -> list[dict]:
    """
    Pronađi sve blokove (``` ili ```mermaid) i za svaki vrati indeks, kod i caption (sljedeći paragraf nakon ```).
    only_mermaid=True: samo blokovi koji sadrže "graph " (Mermaid). only_mermaid=False: svi blokovi (za 804 pozicija).
    Vraća listu dict: [{"index": 1, "caption": "...", "code": "..."}, ...].
    """
    fence = "```"
    start_pattern = re.compile(r"```(?:mermaid)?\s*\n", re.IGNORECASE)
    # Pozicije početaka paragrafa u full_text (redom)
    offsets = []
    pos = 0
    for p in paragraphs:
        offsets.append(pos)
        pos += len(p) + 1
    result = []
    abs_offset = 0  # apsolutni indeks u full_text
    text = full_text
    block_index = 0
    while True:
        match = start_pattern.search(text)
        if not match:
            break
        start_end = match.end()
        rest = text[start_end:]
        close_plain = rest.find("\n" + fence)
        close_with_end = re.search(r"\nend\s*\n\s*" + re.escape(fence), rest, re.IGNORECASE)
        if close_with_end:
            end_pos = close_with_end.start()
            code = rest[:end_pos].strip()
            if code.endswith("end"):
                code = code[:-3].strip()
            close_pos_in_rest = close_with_end.end()
        elif close_plain != -1:
            code = rest[:close_plain].strip()
            close_pos_in_rest = close_plain + len("\n" + fence)
        else:
            code = rest.strip()
            close_pos_in_rest = len(rest)
        code = code.strip()
        is_mermaid = bool(code and (code.startswith("graph ") or (code.split("\n")[0].strip().startswith("%%") and "graph " in code)))
        if only_mermaid and not is_mermaid:
            pass
        else:
            block_index += 1
            # Apsolutna pozicija kraja bloka (nakon ```) u full_text
            end_abs = abs_offset + start_end + close_pos_in_rest
            block_start = abs_offset + match.start()
            # Sljedeći paragraf = prvi paragraf čiji početak je nakon end_abs
            caption = ""
            for i, off in enumerate(offsets):
                if off > end_abs and i < len(paragraphs):
                    caption = (paragraphs[i] or "").strip()
                    break
            # Ako je caption nevaljan (kod, prazan, prekratak), uzmi prethodni paragraf
            def is_bogus(c: str) -> bool:
                c = (c or "").strip()
                if not c or len(c) < 4:
                    return True
                if c.startswith("```") or c.lower().startswith("graph ") or c == "graph":
                    return True
                return False
            if is_bogus(caption):
                prev_i = -1
                for i in range(len(offsets)):
                    end_i = offsets[i] + len(paragraphs[i])
                    if end_i <= block_start:
                        prev_i = i
                if prev_i >= 0:
                    caption = (paragraphs[prev_i] or "").strip()
            result.append({"index": block_index, "caption": clean_caption(caption), "code": code})
        next_close = rest.find(fence)
        if next_close == -1:
            break
        abs_offset += start_end + next_close + len(fence)
        text = rest[next_close + len(fence) :]
    return result


def write_mermaid_diagrami_md(blocks: list[str], out_path: str) -> None:
    """Napiši katalog u mermaid_diagrami.md."""
    lines = ["# Katalog Mermaid dijagrama iz DOCX rukopisa", ""]
    for i, code in enumerate(blocks, 1):
        lines.append(f"## Dijagram {i}")
        lines.append("")
        lines.append("```mermaid")
        lines.append(code)
        lines.append("```")
        lines.append("")
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def main() -> int:
    import argparse
    script_dir = Path(__file__).resolve().parent
    project_root = script_dir.parent
    default_docx = project_root / "manuscript" / "BenediktPerak_knjiga_book_Razvoj_LLM_i_komunikacijskih_agenata_version_02_4_final_2_processed.docx"
    parser = argparse.ArgumentParser(description="Izvlačenje Mermaid blokova i captiona iz DOCX.")
    parser.add_argument("docx", nargs="?", default=None, help="Putanja do DOCX (default: manuscript/...processed.docx)")
    parser.add_argument("--count", action="store_true", help="Samo ispiši broj svih ``` blokova i blokova s 'graph '")
    parser.add_argument("--all-blocks", action="store_true", help="Izvuci SVE ``` blokove (i bez 'graph ') za mapiranje 804 pozicija")
    args = parser.parse_args()

    if args.docx:
        docx_path = (project_root / args.docx).resolve() if not Path(args.docx).is_absolute() else Path(args.docx).resolve()
    else:
        docx_path = default_docx

    if not docx_path.is_file():
        print("Datoteka ne postoji:", docx_path.name, file=sys.stderr)
        return 1

    paragraphs = get_paragraphs_from_docx(docx_path)
    full_text = "\n".join(paragraphs)

    if args.count:
        total, with_graph = count_all_fenced_blocks(full_text)
        print("Ukupno ``` blokova:", total)
        print("Blokova s 'graph ' (Mermaid):", with_graph)
        return 0

    only_mermaid = not args.all_blocks
    blocks_with_captions = extract_blocks_with_captions(full_text, paragraphs, only_mermaid=only_mermaid)
    # Za katalog (mermaid_diagrami.md) pišemo samo blokove s "graph " da ostane valjan
    def is_mermaid_code(code: str) -> bool:
        c = (code or "").strip()
        return bool(c and (c.startswith("graph ") or (c.split("\n")[0].strip().startswith("%%") and "graph " in c)))
    blocks_for_catalog = [b["code"] for b in blocks_with_captions if is_mermaid_code(b["code"])]
    if only_mermaid:
        print("Izvukao sam", len(blocks_with_captions), "Mermaid blokova (samo s 'graph ').")
    else:
        print("Izvukao sam", len(blocks_with_captions), "blokova (svi ```). Caption lista za mapiranje:", len(blocks_with_captions))

    out_md = project_root / "manuscript" / "mermaid_diagrami.md"
    write_mermaid_diagrami_md(blocks_for_catalog, str(out_md))
    print("Katalog zapisan: manuscript/" + out_md.name + " (" + str(len(blocks_for_catalog)) + " blokova s graph)")

    captions_list = [{"index": b["index"], "caption": b["caption"]} for b in blocks_with_captions]
    captions_path = project_root / "manuscript" / "diagram_captions_from_docx.json"
    captions_path.parent.mkdir(parents=True, exist_ok=True)
    with open(captions_path, "w", encoding="utf-8") as f:
        json.dump(captions_list, f, ensure_ascii=False, indent=2)
    print("Caption lista zapisana: manuscript/" + captions_path.name)

    return 0


if __name__ == "__main__":
    sys.exit(main())
