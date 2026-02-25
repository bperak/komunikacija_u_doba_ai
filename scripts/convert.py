from docx import Document
import os
from docx.shared import Cm
import base64
from docx.oxml.shared import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table

def get_image_format(stream):
    """Određuje format slike iz stream-a."""
    if stream.startswith(b'\x89PNG\r\n'):
        return 'png'
    elif stream.startswith(b'\xff\xd8'):
        return 'jpg'
    elif stream.startswith(b'GIF87a') or stream.startswith(b'GIF89a'):
        return 'gif'
    return 'png'  # default format

def handle_run_formatting(run):
    """Obrađuje formatiranje teksta (bold, italic, itd.)"""
    text = run.text
    if text.strip():
        if run.bold and run.italic:
            return f"***{text}***"
        elif run.bold:
            return f"**{text}**"
        elif run.italic:
            return f"*{text}*"
        elif run.underline:
            return f"_{text}_"
        elif run.style and "Hyperlink" in run.style.name:
            # Pronađi hyperlink
            for rel in run.part.rels:
                if run.part.rels[rel].target_ref:
                    return f"[{text}]({run.part.rels[rel].target_ref})"
        return text
    return ""

def convert_table_to_markdown(table):
    """Pretvara tablicu u markdown format."""
    markdown = "\n"
    max_cols = max(len(row.cells) for row in table.rows)
    
    # Zaglavlje tablice
    for cell in table.rows[0].cells:
        markdown += f"| {cell.text.strip()} "
    markdown += "|\n"
    
    # Separator
    markdown += "|" + "---|" * max_cols + "\n"
    
    # Sadržaj tablice
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:  # Preskoči zaglavlje
            continue
        for cell in row.cells:
            markdown += f"| {cell.text.strip()} "
        markdown += "|\n"
    
    return markdown + "\n"

def handle_list_paragraph(paragraph):
    """Obrađuje liste (numerirane i nenumerirane)."""
    text = ""
    for run in paragraph.runs:
        text += handle_run_formatting(run)
    
    # Provjeri stil paragrafa za liste
    if paragraph.style.name.startswith('List Bullet'):
        return f"* {text}\n"
    elif paragraph.style.name.startswith('List Number'):
        return f"1. {text}\n"
    return text + "\n"

def convert_paragraph_to_markdown(paragraph):
    """Pretvara paragraf u markdown format."""
    # Provjera stila paragrafa za naslove
    if paragraph.style.name.startswith('Heading 1'):
        return f"# {paragraph.text}\n\n"
    elif paragraph.style.name.startswith('Heading 2'):
        return f"## {paragraph.text}\n\n"
    elif paragraph.style.name.startswith('Heading 3'):
        return f"### {paragraph.text}\n\n"
    
    # Obrada normalnog teksta s formatiranjem
    text = ""
    for run in paragraph.runs:
        content = run.text
        # Dodaj markdown formatiranje
        if run.bold and run.italic:
            content = f"***{content}***"
        elif run.bold:
            content = f"**{content}**"
        elif run.italic:
            content = f"*{content}*"
        text += content
    
    return text + "\n\n" if text.strip() else "\n"

def word_to_markdown(docx_path):
    """Pretvara Word dokument u Markdown."""
    try:
        # Učitaj Word dokument
        doc = Document(docx_path)
        markdown_content = ""
        
        # Konvertiranje svakog paragrafa
        for paragraph in doc.paragraphs:
            markdown_content += convert_paragraph_to_markdown(paragraph)
        
        # Stvaranje naziva za markdown datoteku
        markdown_path = os.path.splitext(docx_path)[0] + '.md'
        
        # Spremanje markdown sadržaja
        with open(markdown_path, 'w', encoding='utf-8') as md_file:
            md_file.write(markdown_content)
        
        return markdown_path
    except Exception as e:
        print(f"Greška pri obradi dokumenta {docx_path}: {str(e)}")
        raise

def main():
    import sys
    import io
    # Na Windowsu omogući UTF-8 za print (putanje s č, š, ž)
    if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    # Projekt root = jedan nivo iznad scripts/
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    default_docx = os.path.join(
        project_root, "manuscript",
        "Ben knjiga lektorirana_za_obradu.docx"
    )
    target_file = sys.argv[1] if len(sys.argv) > 1 else default_docx
    if not os.path.isabs(target_file):
        target_file = os.path.normpath(os.path.join(project_root, target_file))
    print(f"Trenutni direktorij: {project_root}")
    if not os.path.isfile(target_file):
        print(f"Datoteka '{target_file}' ne postoji!")
        return
    try:
        print(f"Konvertiram: {target_file}")
        doc = Document(target_file)
        markdown_content = ""
        for paragraph in doc.paragraphs:
            markdown_content += convert_paragraph_to_markdown(paragraph)
        markdown_path = os.path.splitext(target_file)[0] + '.md'
        with open(markdown_path, 'w', encoding='utf-8') as md_file:
            md_file.write(markdown_content)
        print(f"Uspješno konvertiran u: {markdown_path}")
    except Exception as e:
        print(f"Greška pri konverziji: {str(e)}")
        import traceback
        print(traceback.format_exc())

if __name__ == "__main__":
    main() 